"""generate_report worker — renders the analysis as a DOCX a person can send.

Two things this worker must never do, both of which it used to:

* Report success without producing a file. A missing library took the "ready"
  path with no ``storage_path``, so the export listed as ready and the download
  answered 202 forever, which reads to a user as "I cannot download my report".
* Leave a report in ``generating``. Any failure now lands on ``failed``, which
  the workspace can show and a person can retry.

A report is available from the moment a read finishes. Waiting for a bid
decision to export the reading is backwards — the reading is what the decision
is made from.
"""

from __future__ import annotations

import os
from datetime import UTC, datetime

from sqlalchemy import select

from app.core.config import get_settings
from app.core.logging import get_logger
from app.db.base import async_session_factory
from app.db.models.analysis import Analysis
from app.db.models.requirement import Requirement
from app.db.models.question import Question
from app.db.models.report import Report

logger = get_logger()

FINDING_SECTIONS = (
    ("Identity", "identity"),
    ("Scope", "scope"),
    ("Legal and regulatory", "legal"),
    ("Eligibility", "eligibility"),
    ("Pricing", "pricing"),
    ("Post-award", "post_award"),
)

DECISION_LABELS = {
    "bid": "Bid",
    "no-bid": "No-bid",
    "watch": "Watch",
    "undecided": "Not yet recorded",
}

STAGE_LABELS = {
    "triage": "Triage",
    "analyzing": "Reading in progress",
    "review": "In review",
    "decided": "Decided",
}


async def generate_report_task(ctx: dict, report_id: str) -> dict:
    """Arq task: render a DOCX report from everything the analysis holds."""
    settings = get_settings()
    redis = ctx.get("redis", ctx.get("job_ctx", {}).get("redis"))

    async with async_session_factory() as db:
        result = await db.execute(select(Report).where(Report.id == report_id))
        report = result.scalar_one_or_none()
        if not report:
            logger.error("report_not_found", report_id=report_id)
            return {"error": "Report not found"}

        try:
            analysis_result = await db.execute(
                select(Analysis).where(Analysis.id == report.analysis_id)
            )
            analysis = analysis_result.scalar_one_or_none()
            if not analysis:
                return await _fail(db, report, "Analysis not found")

            rows = (
                await db.execute(
                    select(Requirement)
                    .where(Requirement.analysis_id == analysis.id, Requirement.state != "removed")
                    .order_by(Requirement.document_id, Requirement.page, Requirement.reference)
                )
            ).scalars().all()
            questions = (
                await db.execute(
                    select(Question)
                    .where(Question.analysis_id == analysis.id)
                    .order_by(Question.order.asc())
                )
            ).scalars().all()

            logger.info(
                "generate_report_start",
                report_id=report_id,
                analysis_id=analysis.id,
                format=report.format,
            )
            blocks = None
            if EVIDENCE_PACK in (report.template_name or "").lower():
                blocks = await _evidence_blocks(db, analysis, rows)
            filepath = render(settings.REPORTS_DIR, report, analysis, rows, questions, blocks)

            report.status = "ready"
            report.storage_path = filepath
            report.size = os.path.getsize(filepath)
            await db.commit()

            if redis:
                import orjson

                await redis.publish(
                    f"notifications:{report.org_id}",
                    orjson.dumps({"event": "report_ready", "reportId": report_id}).decode(),
                )

            logger.info("generate_report_complete", report_id=report_id, bytes=report.size)
            return {"status": "completed", "path": filepath}

        except Exception as exc:  # noqa: BLE001 — the record must never be left mid-flight
            logger.exception("generate_report_error", report_id=report_id, error=str(exc))
            return await _fail(db, report, str(exc)[:300])



async def _evidence_blocks(db, analysis: Analysis, rows: list[Requirement]) -> list:
    """Assemble the evidence pack from what is stored.

    The verification queue is rebuilt here rather than read from a table on
    purpose: it is derived from the current state of everything else, and a
    stored copy would be a second version of the truth that quietly drifts.
    """
    from app.db.models.response_check import ResponseCheck
    from app.pipeline import verification
    from app.reports import evidence

    version = int((analysis.response or {}).get("version") or 0)
    checks: list = []
    if version:
        checks = list(
            (
                await db.execute(
                    select(ResponseCheck).where(
                        ResponseCheck.analysis_id == analysis.id,
                        ResponseCheck.response_version == version,
                    )
                )
            )
            .scalars()
            .all()
        )
    from app.db.models.question import Question

    questions = list(
        (await db.execute(select(Question).where(Question.analysis_id == analysis.id)))
        .scalars()
        .all()
    )
    from app.db.models.review import ReviewFinding, ReviewRound

    rounds = list(
        (await db.execute(select(ReviewRound).where(ReviewRound.analysis_id == analysis.id)))
        .scalars()
        .all()
    )
    review_findings = list(
        (await db.execute(select(ReviewFinding).where(ReviewFinding.analysis_id == analysis.id)))
        .scalars()
        .all()
    )
    from app.db.models.contradiction import Contradiction as ContradictionRow

    conflicts = list(
        (await db.execute(select(ContradictionRow).where(ContradictionRow.analysis_id == analysis.id)))
        .scalars()
        .all()
    )
    queue = verification.build(
        analysis=analysis,
        requirements=list(rows),
        checks=checks,
        questions=questions,
        reviews=rounds,
        review_findings=review_findings,
        contradictions=conflicts,
    )
    return evidence.build(
        analysis=analysis,
        requirements=list(rows),
        checks=checks,
        queue=queue,
        questions=questions,
        reviews=rounds,
        review_findings=review_findings,
        contradictions=conflicts,
    )


async def _fail(db, report: Report, reason: str) -> dict:
    report.status = "failed"
    await db.commit()
    return {"error": reason}


#: The template name that produces the evidence pack rather than a briefing.
EVIDENCE_PACK = "evidence pack"


def render(
    reports_dir: str,
    report: Report,
    analysis: Analysis,
    rows: list[Requirement],
    questions: list[Question],
    blocks: list | None = None,
) -> str:
    """Render in the format that was asked for.

    Naming a DOCX ``.pdf`` is worse than not offering PDF: the file opens in
    nothing and the user has no idea why. So PDF is a real conversion or a real
    failure, never a rename.

    ``blocks`` carries the evidence pack when one was requested. It is passed
    in rather than assembled here so both formats render the same record — a
    pack that differed between DOCX and Markdown would be two records.
    """
    fmt = (report.format or "DOCX").upper()
    if blocks is not None:
        if fmt == "MD":
            return _blocks_markdown(reports_dir, report, blocks)
        docx_path = _blocks_docx(reports_dir, report, blocks)
        return docx_path if fmt != "PDF" else _to_pdf(docx_path)

    if fmt == "MD":
        return _render_markdown(reports_dir, report, analysis, rows, questions)

    docx_path = _render_docx(reports_dir, report, analysis, rows, questions)
    if fmt != "PDF":
        return docx_path
    return _to_pdf(docx_path)



# ── Evidence pack ────────────────────────────────────────────────────────
#
# The pack arrives as blocks — headings, paragraphs, notes and tables — so the
# two formats walk identical content. Anything a renderer decided for itself
# would be a difference between two copies of the same record.


def _blocks_docx(reports_dir: str, report: Report, blocks: list) -> str:
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    for block in blocks:
        kind = block[0]
        if kind == "heading":
            doc.add_heading(str(block[2]), level=int(block[1]))
        elif kind == "para":
            doc.add_paragraph(str(block[1]))
        elif kind == "note":
            paragraph = doc.add_paragraph(str(block[1]))
            paragraph.runs[0].font.size = Pt(8)
            paragraph.runs[0].italic = True
        elif kind == "table":
            headers, rows = block[1], block[2]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            for index, name in enumerate(headers):
                table.rows[0].cells[index].text = str(name)
            for row in rows:
                cells = table.add_row().cells
                for index, value in enumerate(row[: len(headers)]):
                    cells[index].text = str(value)
            doc.add_paragraph()

    os.makedirs(reports_dir, exist_ok=True)
    filepath = os.path.join(reports_dir, f"{report.id}.docx")
    doc.save(filepath)
    return filepath


def _blocks_markdown(reports_dir: str, report: Report, blocks: list) -> str:
    out: list[str] = []
    for block in blocks:
        kind = block[0]
        if kind == "heading":
            out += ["", f"{'#' * int(block[1])} {block[2]}", ""]
        elif kind == "para":
            out += [str(block[1]), ""]
        elif kind == "note":
            out += [f"*{block[1]}*", ""]
        elif kind == "table":
            headers, rows = block[1], block[2]
            out.append("| " + " | ".join(str(h) for h in headers) + " |")
            out.append("| " + " | ".join("---" for _ in headers) + " |")
            for row in rows:
                cells = [str(value).replace("|", "\\|").replace("\n", " ") for value in row[: len(headers)]]
                cells += [""] * (len(headers) - len(cells))
                out.append("| " + " | ".join(cells) + " |")
            out.append("")

    os.makedirs(reports_dir, exist_ok=True)
    filepath = os.path.join(reports_dir, f"{report.id}.md")
    with open(filepath, "w", encoding="utf-8") as handle:
        handle.write("\n".join(out).strip() + "\n")
    return filepath


def _render_docx(
    reports_dir: str,
    report: Report,
    analysis: Analysis,
    rows: list[Requirement],
    questions: list[Question],
) -> str:
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()

    doc.add_heading(analysis.title, 0)
    doc.add_paragraph(f"{report.template_name} · generated {datetime.now(UTC).strftime('%d %B %Y, %H:%M UTC')}")

    facts = doc.add_paragraph()
    for label, value in (
        ("Agency", analysis.agency),
        ("Solicitation", analysis.solicitation_number),
        ("Document type", (analysis.doc_type or "").upper()),
        ("Stage", STAGE_LABELS.get(analysis.stage, analysis.stage)),
        ("Decision", DECISION_LABELS.get(analysis.go_no_go, analysis.go_no_go)),
    ):
        if not value:
            continue
        run = facts.add_run(f"{label}: ")
        run.bold = True
        facts.add_run(f"{value}\n")

    if analysis.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(analysis.summary)

    _decision(doc, analysis)
    _calendar(doc, analysis)
    _research(doc, analysis)
    _evaluation(doc, analysis)
    _risks(doc, analysis)
    _findings(doc, analysis)
    _matrix(doc, rows)
    _questions(doc, questions)

    note = doc.add_paragraph(
        "Every citation in this report was resolved against the uploaded document. "
        "Findings whose quote could not be located in the extract are marked "
        "“source not located” and need a human check."
    )
    note.runs[0].font.size = Pt(8)

    os.makedirs(reports_dir, exist_ok=True)
    filepath = os.path.join(reports_dir, f"{report.id}.docx")
    doc.save(filepath)
    return filepath


def _to_pdf(docx_path: str) -> str:
    """Convert through LibreOffice when the image carries it."""
    import shutil
    import subprocess
    import tempfile

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError(
            "PDF export needs LibreOffice, which is not installed in this image. "
            "Export as DOCX or Markdown, or add libreoffice-writer to the worker image."
        )
    outdir = os.path.dirname(docx_path)
    with tempfile.TemporaryDirectory(prefix="soffice-") as profile:
        subprocess.run(  # noqa: S603 — fixed binary, paths we produced ourselves
            [
                soffice,
                # LibreOffice insists on a writable profile directory. Pointing
                # it at a temp one keeps the worker container's filesystem
                # read-only, and keeps concurrent conversions from colliding.
                f"-env:UserInstallation=file://{profile}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                outdir,
                docx_path,
            ],
            check=True,
            capture_output=True,
            timeout=180,
        )
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    if not os.path.exists(pdf_path):
        raise RuntimeError("LibreOffice reported success but produced no PDF.")
    os.remove(docx_path)
    return pdf_path


def _render_markdown(
    reports_dir: str,
    report: Report,
    analysis: Analysis,
    rows: list[Requirement],
    questions: list[Question],
) -> str:
    """The same report as plain text, for a wiki, an email, or a diff."""
    out: list[str] = [f"# {analysis.title}", ""]
    out.append(f"*{report.template_name} · generated {datetime.now(UTC).strftime('%d %B %Y, %H:%M UTC')}*")
    out.append("")
    for label, value in (
        ("Agency", analysis.agency),
        ("Solicitation", analysis.solicitation_number),
        ("Document type", (analysis.doc_type or "").upper()),
        ("Stage", STAGE_LABELS.get(analysis.stage, analysis.stage)),
        ("Decision", DECISION_LABELS.get(analysis.go_no_go, analysis.go_no_go)),
    ):
        if value:
            out.append(f"- **{label}:** {value}")
    if analysis.summary:
        out += ["", "## Summary", "", analysis.summary]

    gates = analysis.gates or []
    if gates:
        out += ["", "## Go / no-go", ""]
        for gate in gates:
            answer = {True: "Met", False: "Not met", None: "Unresolved"}[gate.get("met")]
            out.append(f"- **[{answer}]** {gate.get('question', '')} — {gate.get('answer', '')}")
            quoted = _cite(gate.get("citation"))
            if quoted:
                out.append(f"  > {quoted}")

    if analysis.dates:
        out += ["", "## Key dates", "", "| Date | Milestone | Source |", "| --- | --- | --- |"]
        for date in analysis.dates:
            source = "Stated in document" if date.get("source") != "derived" else "Planned by Margin"
            out.append(f"| {str(date.get('at', ''))[:16].replace('T', ' ')} | {date.get('label', '')} | {source} |")

    research = analysis.research or {}
    if str(research.get("status") or "not_requested") != "not_requested":
        out += ["", "## External research", ""]
        out.append(
            "*Read on the open web, not in the solicitation. No clause reference, because there is none to give.*"
        )
        if str(research.get("status")) != "completed":
            out.append("")
            out.append(f"This pass returned no research ({str(research.get('status')).replace('_', ' ')}).")
        else:
            if research.get("query"):
                out += ["", f"**Searched for:** {research['query']}"]
            sources = research.get("sources") or []
            by_url = {str(s.get("url")): s for s in sources}
            claims = research.get("claims") or []
            if claims:
                for claim in claims:
                    out += ["", str(claim.get("text") or "")]
                    cited = [by_url.get(url) for url in (claim.get("sources") or [])]
                    links = [
                        f"[{c.get('site') or c.get('title')}]({c.get('url')})" for c in cited if c
                    ]
                    out.append(
                        f"  *Source: {', '.join(links)}*" if links
                        else "  *No source was cited for this paragraph.*"
                    )
            elif research.get("summary"):
                out += ["", str(research["summary"])]
            out += ["", "### Sources", ""]
            if not sources:
                out.append("This pass returned no sources. Treat the above as a lead to check, not as evidence.")
            for source in sources:
                title = str(source.get("title") or source.get("url") or "")
                out.append(f"- [{title}]({source.get('url', '')}) — {source.get('site', '')}")

    for section_name, attr in FINDING_SECTIONS:
        findings = getattr(analysis, attr) or []
        if not findings:
            continue
        out += ["", f"## {section_name}", ""]
        for finding in findings:
            out.append(f"- **{finding.get('label', '')}:** {finding.get('value', '')}")
            quoted = _cite(finding.get("citation"))
            if quoted:
                out.append(f"  > {quoted}")

    if rows:
        out += ["", "## Compliance matrix", "", "| Reference | Requirement | Type | Owner | Status |", "| --- | --- | --- | --- | --- |"]
        for row in rows:
            requirement = (row.text or "").replace("|", "\\|")
            out.append(
                f"| {row.reference or ''} | {requirement} | {row.type or ''} | "
                f"{row.owner or 'Unassigned'} | {row.status or ''} |"
            )

    if questions:
        out += ["", "## Questions for the agency", ""]
        for index, question in enumerate(questions, start=1):
            out.append(f"{index}. {question.text or ''}")
            if question.rationale:
                out.append(f"   *{question.rationale}*")

    os.makedirs(reports_dir, exist_ok=True)
    filepath = os.path.join(reports_dir, f"{report.id}.md")
    with open(filepath, "w", encoding="utf-8") as handle:
        handle.write("\n".join(out) + "\n")
    return filepath


def _small(paragraph) -> None:  # noqa: ANN001 — python-docx has no stubs
    from docx.shared import Pt

    for run in paragraph.runs:
        run.font.size = Pt(8)
        run.italic = True


def _cite(citation: dict | None) -> str:
    if not isinstance(citation, dict) or not citation.get("quote"):
        return ""
    if citation.get("located") is False:
        return f'“{citation["quote"]}” (source not located)'
    page = citation.get("page")
    section = str(citation.get("section") or "").strip()
    where = f"p.{page}" if page else ""
    if section:
        where = f"{where} · {section}" if where else section
    return f'“{citation["quote"]}”' + (f" ({where})" if where else "")


def _quote(doc, citation: dict | None) -> None:
    text = _cite(citation)
    if not text:
        return
    try:
        doc.add_paragraph(text, style="Quote")
    except KeyError:
        # Some base templates have no Quote style; the text matters, not the style.
        doc.add_paragraph(text)


def _decision(doc, analysis: Analysis) -> None:
    gates = analysis.gates or []
    if not gates:
        return
    doc.add_heading("Go / no-go", level=1)
    met = sum(1 for g in gates if g.get("met") is True)
    hard = sum(1 for g in gates if g.get("weight") == "hard")
    doc.add_paragraph(
        f"{met} of {len(gates)} gates met · {hard} hard "
        f"{'gate' if hard == 1 else 'gates'} · decision: "
        f"{DECISION_LABELS.get(analysis.go_no_go, analysis.go_no_go)}"
    )
    if analysis.decision_note:
        doc.add_paragraph(analysis.decision_note)
    for gate in gates:
        answer = {True: "Met", False: "Not met", None: "Unresolved"}[gate.get("met")]
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(f"[{answer}] ").bold = True
        para.add_run(f"{gate.get('question', '')} — {gate.get('answer', '')}")
        _quote(doc, gate.get("citation"))


def _calendar(doc, analysis: Analysis) -> None:
    dates = analysis.dates or []
    if not dates:
        return
    doc.add_heading("Key dates", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Date"
    header[1].text = "Milestone"
    header[2].text = "Source"
    for date in dates:
        cells = table.add_row().cells
        cells[0].text = str(date.get("at", ""))[:16].replace("T", " ")
        cells[1].text = str(date.get("label", ""))
        cells[2].text = "Stated in document" if date.get("source") != "derived" else "Planned by Margin"


def _research(doc, analysis: Analysis) -> None:
    """External research, under its own heading and with its sources listed.

    A reviewer reading this document offline has no way to ask where a claim
    came from, so the answer has to be printed next to it.
    """
    research = analysis.research or {}
    status = str(research.get("status") or "not_requested")
    if status == "not_requested":
        return

    doc.add_heading("External research", level=1)
    doc.add_paragraph(
        "The following was read on the open web, not in the solicitation. It is "
        "context for the rules this procurement sits inside, and it carries no "
        "clause reference because there is none to give."
    )
    if status != "completed":
        doc.add_paragraph(
            f"This pass returned no research ({status.replace('_', ' ')}). "
            "Every other section of this report comes from the document itself."
        )
        return

    if research.get("query"):
        doc.add_paragraph(f"Searched for: {research['query']}")

    sources = research.get("sources") or []
    by_url = {str(s.get("url")): s for s in sources}
    claims = research.get("claims") or []
    if claims:
        # Each paragraph is followed by the pages that back it. A reader
        # checking one sentence should not have to guess which of a dozen
        # sources at the end of the section it came from.
        for claim in claims:
            doc.add_paragraph(str(claim.get("text") or ""))
            cited = [by_url.get(url) for url in (claim.get("sources") or [])]
            named = [str(c.get("site") or c.get("title") or "") for c in cited if c]
            attribution = doc.add_paragraph(
                f"Source: {', '.join(named)}" if named
                else "No source was cited for this paragraph — check it before relying on it."
            )
            _small(attribution)
    elif research.get("summary"):
        for para in str(research["summary"]).split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    if not sources:
        doc.add_paragraph(
            "This pass returned no sources, so nothing above can be traced to a "
            "page. Treat it as a lead to check, not as evidence."
        )
        return
    doc.add_heading("Sources", level=2)
    for source in sources:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(str(source.get("title") or source.get("url") or "")).bold = True
        site = str(source.get("site") or "")
        para.add_run(f" — {site}" if site else "")
        doc.add_paragraph(str(source.get("url") or ""))


def _evaluation(doc, analysis: Analysis) -> None:
    factors = analysis.evaluation or []
    if not factors:
        return
    doc.add_heading("Evaluation factors", level=1)
    for factor in factors:
        para = doc.add_paragraph(style="List Bullet")
        weight = factor.get("weight") or 0
        para.add_run(f"{factor.get('name', '')}").bold = True
        para.add_run(f" — {weight}%" if weight else "")
        if factor.get("method"):
            para.add_run(f": {factor['method']}")
        _quote(doc, factor.get("citation"))


def _risks(doc, analysis: Analysis) -> None:
    risks = analysis.risks or []
    if not risks:
        return
    doc.add_heading("Risks", level=1)
    for risk in risks:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(f"{risk.get('title', '')} [{risk.get('severity', '')}]").bold = True
        if risk.get("narrative"):
            para.add_run(f" — {risk['narrative']}")
        if risk.get("mitigation"):
            doc.add_paragraph(f"Mitigation: {risk['mitigation']}")
        _quote(doc, risk.get("citation"))


def _findings(doc, analysis: Analysis) -> None:
    for section_name, attr in FINDING_SECTIONS:
        findings = getattr(analysis, attr) or []
        if not findings:
            continue
        doc.add_heading(section_name, level=1)
        for finding in findings:
            para = doc.add_paragraph()
            para.add_run(f"{finding.get('label', '')}: ").bold = True
            para.add_run(str(finding.get("value", "")))
            if finding.get("detail"):
                doc.add_paragraph(str(finding["detail"]))
            _quote(doc, finding.get("citation"))


def _matrix(doc, rows: list[Requirement]) -> None:
    if not rows:
        return
    doc.add_heading("Compliance matrix", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for index, name in enumerate(("Reference", "Requirement", "Type", "Owner", "Status")):
        table.rows[0].cells[index].text = name
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row.reference or ""
        cells[1].text = row.text or ""
        cells[2].text = row.type or ""
        cells[3].text = row.owner or "Unassigned"
        cells[4].text = row.status or ""


def _questions(doc, questions: list[Question]) -> None:
    if not questions:
        return
    doc.add_heading("Questions for the agency", level=1)
    for question in questions:
        para = doc.add_paragraph(style="List Number")
        para.add_run(question.text or "")
        if question.rationale:
            doc.add_paragraph(f"Why: {question.rationale}")
        _quote(doc, question.citation)
