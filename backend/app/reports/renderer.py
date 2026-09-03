"""Branded DOCX report renderer using python-docx and Jinja template logic.

Report Structure:
1. Executive Summary & Go/No-Go verdict
2. Key Dates & Logistics
3. Scope Decoded
4. Compliance Matrix (Dynamic table)
5. Legal & Regulatory (In-document and external citations)
6. Eligibility & Evaluation Factors
7. Risk & Red-Flag Audit
8. Clarifying Questions (Portal ready)
9. Appendix: Complete SILENT Findings Ledger
"""

from __future__ import annotations

import os
from datetime import UTC, datetime
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.core.logging import get_logger

logger = get_logger()

# Brand palette (from design spec):
# Seal (Mandatory / Disqualifying): #9E2A2B (Deep brick/seal)
# Ochre (Scored / Important): #D97706 (Amber/ochre)
# Leaf (Verified / Met): #15803D (Forest green)
# Slate (Muted / Informational): #475569 (Cool slate)
COLOR_SEAL = RGBColor(158, 42, 43)
COLOR_OCHRE = RGBColor(217, 119, 6)
COLOR_LEAF = RGBColor(21, 128, 61)
COLOR_SLATE = RGBColor(71, 85, 105)


def _set_cell_background(cell, fill_color: str):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tc_pr.append(shd)


class ReportRenderer:
    """Generates styled Word (.docx) reports matching the Margin branding guidelines."""

    def render_analysis_report(
        self,
        analysis: dict[str, Any],
        matrix_rows: list[dict[str, Any]],
        questions: list[dict[str, Any]],
        output_path: str,
    ) -> str:
        doc = Document()

        # Set 1-inch margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # ── Document Title & Header ──────────────────────────────────────
        title_p = doc.add_paragraph()
        title_run = title_p.add_run("MARGIN CAPTURE INTELLIGENCE")
        title_run.font.name = "Georgia"
        title_run.font.size = Pt(9)
        title_run.font.bold = True
        title_run.font.color.rgb = COLOR_SLATE

        heading = doc.add_heading(analysis.get("title", "Solicitation Analysis"), level=1)
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after = Pt(8)

        meta_p = doc.add_paragraph()
        meta_p.add_run(f"Agency: {analysis.get('agency', 'N/A')}  |  ").bold = True
        meta_p.add_run(f"Solicitation: {analysis.get('solicitationNumber', 'N/A')}  |  ")
        meta_p.add_run(f"Stage: {analysis.get('stage', 'review').title()}  |  ")
        meta_p.add_run(f"Generated: {datetime.now(UTC).strftime('%B %d, %Y')}")

        # ── 1. Go/No-Go Executive Summary ─────────────────────────────────
        doc.add_heading("1. Executive Summary & Go/No-Go Decision", level=2)
        dec_p = doc.add_paragraph()
        dec_label = dec_p.add_run("RECOMMENDED DECISION: ")
        dec_label.bold = True
        verdict = analysis.get("goNoGo", "undecided").upper()
        verdict_run = dec_p.add_run(verdict)
        verdict_run.bold = True
        if verdict == "BID":
            verdict_run.font.color.rgb = COLOR_LEAF
        elif verdict == "NO-BID":
            verdict_run.font.color.rgb = COLOR_SEAL
        else:
            verdict_run.font.color.rgb = COLOR_OCHRE

        if analysis.get("decisionNote"):
            doc.add_paragraph(f"Decision Note: {analysis['decisionNote']}")

        summary_text = analysis.get("summary") or "Comprehensive capture analysis completed across all evaluation factors."
        doc.add_paragraph(summary_text)

        # ── 2. Key Dates & Logistics ─────────────────────────────────────
        doc.add_heading("2. Key Dates & Logistics", level=2)
        dates = analysis.get("dates", [])
        if dates:
            d_table = doc.add_table(rows=1, cols=3)
            d_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = d_table.rows[0].cells
            hdr_cells[0].text = "Event"
            hdr_cells[1].text = "Due Date / Time"
            hdr_cells[2].text = "Kind"
            for c in hdr_cells:
                _set_cell_background(c, "F1F5F9")
                c.paragraphs[0].runs[0].font.bold = True

            for d in dates:
                row_cells = d_table.add_row().cells
                row_cells[0].text = d.get("label", "")
                row_cells[1].text = f"{d.get('at', '')} ({d.get('timezone', 'UTC')})"
                row_cells[2].text = d.get("kind", "")
        else:
            doc.add_paragraph("No explicit deadline dates recorded.")

        # ── 3. Scope Decoded ─────────────────────────────────────────────
        doc.add_heading("3. Scope Decoded", level=2)
        for f in analysis.get("scope", []):
            self._render_finding_bullet(doc, f)

        # ── 4. Compliance Matrix Table ───────────────────────────────────
        doc.add_heading("4. Compliance Matrix", level=2)
        if matrix_rows:
            m_table = doc.add_table(rows=1, cols=5)
            m_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            m_hdrs = m_table.rows[0].cells
            m_hdrs[0].text = "Ref"
            m_hdrs[1].text = "Requirement"
            m_hdrs[2].text = "Stakes"
            m_hdrs[3].text = "Status"
            m_hdrs[4].text = "Owner"
            for c in m_hdrs:
                _set_cell_background(c, "F1F5F9")
                c.paragraphs[0].runs[0].font.bold = True

            for r in matrix_rows[:30]:  # Limit for brevity in Word
                r_cells = m_table.add_row().cells
                r_cells[0].text = r.get("reference", "")
                r_cells[1].text = r.get("requirement", "")[:120]
                stakes_run = r_cells[2].paragraphs[0].add_run(r.get("stakes", ""))
                if r.get("stakes") == "disqualifying":
                    stakes_run.font.color.rgb = COLOR_SEAL
                r_cells[3].text = r.get("status", "")
                r_cells[4].text = r.get("owner") or "Unassigned"
        else:
            doc.add_paragraph("Matrix is being drafted.")

        # ── 5. Legal & Regulatory ────────────────────────────────────────
        doc.add_heading("5. Legal & Regulatory Obligations", level=2)
        for f in analysis.get("legal", []):
            self._render_finding_bullet(doc, f)

        # ── 6. Eligibility & Evaluation ──────────────────────────────────
        doc.add_heading("6. Eligibility Gates & Evaluation Criteria", level=2)
        for f in analysis.get("eligibility", []):
            self._render_finding_bullet(doc, f)
        for f in analysis.get("evaluation", []):
            self._render_finding_bullet(doc, f)

        # ── 7. Risks & Red Flags ─────────────────────────────────────────
        doc.add_heading("7. Risk & Red-Flag Audit", level=2)
        for f in analysis.get("risks", []):
            self._render_finding_bullet(doc, f)

        # ── 8. Clarifying Questions ──────────────────────────────────────
        doc.add_heading("8. Clarifying Questions (Portal Ready)", level=2)
        if questions:
            for idx, q in enumerate(questions, 1):
                qp = doc.add_paragraph()
                qp.add_run(f"Q{idx}. {q.get('text', '')}").bold = True
                doc.add_paragraph(f"Rationale: {q.get('rationale', '')}", style="Quote")
        else:
            doc.add_paragraph("No outstanding questions generated.")

        # ── 9. Appendix: SILENT Ledger ───────────────────────────────────
        doc.add_heading("Appendix: SILENT Findings Ledger", level=2)
        doc.add_paragraph("Items where the solicitation is explicitly silent or omits customary standards:")
        silent_items = analysis.get("silent", [])
        if silent_items:
            for s in silent_items:
                sp = doc.add_paragraph(style="List Bullet")
                sp.add_run(f"{s.get('topic', 'Topic')}: ").bold = True
                sp.add_run(f"Expected: {s.get('expectation', '')}. Consequence: {s.get('consequence', '')}")
        else:
            doc.add_paragraph("No omissions registered in the silent ledger.")

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        logger.info("report_rendered", path=output_path)
        return output_path

    def _render_finding_bullet(self, doc: Document, finding: dict[str, Any]) -> None:
        p = doc.add_paragraph(style="List Bullet")
        label_run = p.add_run(f"{finding.get('label', '')}: ")
        label_run.bold = True

        stakes = finding.get("stakes", "informational")
        if stakes == "disqualifying":
            label_run.font.color.rgb = COLOR_SEAL
        elif stakes == "scored":
            label_run.font.color.rgb = COLOR_OCHRE

        p.add_run(str(finding.get("value", "")))

        # Verified badge
        if finding.get("verified"):
            v_run = p.add_run(" [Verified]")
            v_run.font.color.rgb = COLOR_LEAF
            v_run.font.size = Pt(8.5)

        # Inline citation
        citation = finding.get("citation") or {}
        if citation.get("quote"):
            cite_p = doc.add_paragraph(style="Quote")
            cite_run = cite_p.add_run(
                f"Source: \"{citation['quote']}\" (Page {citation.get('page', '?')}, {citation.get('section', 'Section')})"
            )
            cite_run.font.size = Pt(9)
            cite_run.font.color.rgb = COLOR_SLATE
