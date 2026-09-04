"""Text extraction for uploaded documents, one page at a time.

Page boundaries are the whole point of this module. A citation that says
"p.24" is only worth anything if page 24 of the extract is page 24 of the
document a person is holding, so extraction keeps the page breaks the format
gives us instead of flattening everything into one blob and guessing later.

Pages are joined with a form feed (``\\f``) — the character that has meant
"page break" in plain text since line printers — so a single ``raw_text``
column round-trips the structure without a second table.
"""

from __future__ import annotations

from app.core.logging import get_logger

logger = get_logger()

PAGE_SEP = "\f"

TEXT_SUFFIXES = (".txt", ".md", ".markdown", ".csv", ".json", ".xml", ".html", ".htm", ".rtf")

# Roughly a printed page of a US-letter solicitation. Only used for formats
# that genuinely have no page breaks of their own.
LINES_PER_SYNTHETIC_PAGE = 46


def extract_pages(content: bytes, filename: str) -> list[str]:
    """Best-effort page-by-page text. Never raises — an unreadable upload still
    produces an analysis, it just has less to say."""
    lower = filename.lower()

    if lower.endswith(".pdf"):
        pages = _from_pdf(content)
        if pages:
            return pages
    elif lower.endswith(".docx"):
        pages = _from_docx(content)
        if pages:
            return pages

    if lower.endswith(TEXT_SUFFIXES) or not lower.rpartition(".")[2]:
        return _split_plain(content.decode("utf-8", errors="replace"))

    # Unknown binary: decoding is still better than nothing, and the reader
    # downstream drops lines that are not printable.
    decoded = content.decode("utf-8", errors="ignore")
    return _split_plain(decoded) if decoded.strip() else []


def extract_text(content: bytes, filename: str) -> str:
    """The whole document as one string, page breaks preserved as form feeds."""
    return PAGE_SEP.join(extract_pages(content, filename))


def pages_from_text(text: str) -> list[str]:
    """Recover the page list from a stored ``raw_text``.

    Text written before page-aware extraction has no form feeds in it, so it
    falls back to fixed-height paging — wrong, but wrong in a stable way, and
    the anchor still resolves a quote to the page it prints on.
    """
    if PAGE_SEP in text:
        return text.split(PAGE_SEP)
    return _split_plain(text)


def _split_plain(text: str) -> list[str]:
    if PAGE_SEP in text:
        return text.split(PAGE_SEP)
    lines = text.splitlines()
    if not lines:
        return [text] if text else []
    return [
        "\n".join(lines[i : i + LINES_PER_SYNTHETIC_PAGE])
        for i in range(0, len(lines), LINES_PER_SYNTHETIC_PAGE)
    ]


def _from_pdf(content: bytes) -> list[str]:
    try:
        import io

        from pypdf import PdfReader
    except ImportError:
        logger.warning("pdf_extract_unavailable", reason="pypdf not installed")
        return []

    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:  # noqa: BLE001 — a bad PDF must not fail the upload
        logger.warning("pdf_extract_failed", error=str(exc))
        return []

    if not any(pages):
        # A scanned PDF with no text layer. Say so rather than returning a
        # stack of blank pages that would read as a successful extraction.
        logger.warning("pdf_extract_empty", pages=len(pages))
        return []
    return pages


def _from_docx(content: bytes) -> list[str]:
    """DOCX has no fixed pagination, but Word records where it last rendered a
    page break, and an author's explicit breaks are always there. Both are good
    enough to make "p.12" mean something."""
    try:
        import io

        from docx import Document as DocxDocument
    except ImportError:
        return []

    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        logger.warning("docx_extract_failed", error=str(exc))
        return []

    pages: list[list[str]] = [[]]
    try:
        for paragraph in doc.paragraphs:
            if _docx_breaks_page(paragraph) and pages[-1]:
                pages.append([])
            text = paragraph.text.strip()
            if text:
                pages[-1].append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = " | ".join(cell.text.strip() for cell in row.cells)
                if cells.strip(" |"):
                    pages[-1].append(cells)
    except Exception as exc:  # noqa: BLE001
        logger.warning("docx_walk_failed", error=str(exc))

    rendered = ["\n".join(page) for page in pages if page]
    if not rendered:
        return []
    # A DOCX with no break information at all is one giant page; fall back to
    # fixed-height paging so citations still land somewhere specific.
    if len(rendered) == 1:
        return _split_plain(rendered[0])
    return rendered


_PAGE_BREAK_XPATH = (
    './/w:br[@w:type="page"] | .//w:lastRenderedPageBreak'
)


def _docx_breaks_page(paragraph) -> bool:  # noqa: ANN001 — python-docx has no stubs
    try:
        return bool(paragraph._p.xpath(_PAGE_BREAK_XPATH))
    except Exception:  # noqa: BLE001
        return False
